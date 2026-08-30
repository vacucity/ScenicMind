from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "ScenicMind-路演技术手册.md"
DIAGRAM = ROOT / "ScenicMind-Multi-Agent-Architecture.png"
OUTPUT = ROOT / "ScenicMind-路演技术手册.docx"

PINE = "104737"
PINE_LIGHT = "E8F0EA"
AMBER = "C58B35"
INK = "17231D"
MUTED = "68766E"
GRID = "D7E1DB"
SOFT = "F3F0E8"
BLUE = "2E74B5"
BLUE_DARK = "1F4D78"
FONT = "Microsoft YaHei"
MONO = "Consolas"


def set_run_font(run, name=FONT, size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(element, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    if hasattr(element, "get_or_add_pPr"):
        element.get_or_add_pPr().append(shd)
    else:
        element.get_or_add_tcPr().append(shd)


def paragraph_border(paragraph, side="left", color=AMBER, size="18", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), space)
    border.set(qn("w:color"), color)
    p_bdr.append(border)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("SCENICMIND 智景  ·  路演技术手册")
    set_run_font(hr, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fr = fp.add_run("内部路演材料  |  ")
    set_run_font(fr, size=8, color=MUTED)
    add_field(fp, "PAGE")
    for run in fp.runs:
        set_run_font(run, size=8, color=MUTED)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5.5)
    normal.paragraph_format.line_spacing = 1.23

    for style_name, size, color, before, after in [
        ("Heading 1", 17, PINE, 16, 8),
        ("Heading 2", 13.5, PINE, 12, 6),
        ("Heading 3", 11.5, BLUE_DARK, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def create_numbering(doc, ordered):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(existing_abs or [0]) + 1
    num_id = max(existing_num or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "•")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.extend([tabs, ind])
    level.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    p_pr.append(num_pr)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.2


def add_inline(paragraph, text, size=10.5, color=INK):
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    for part in pattern.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, color=PINE, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name=MONO, size=max(8.5, size - 1), color=BLUE_DARK)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{margin}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_table(doc, rows):
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    if cols == 5:
        widths = [2300, 1100, 1300, 1300, 1300]
    elif cols == 4:
        widths = [1550, 2700, 1300, 2250]
    elif cols == 3:
        widths = [1800, 2600, 3400]
    elif cols == 2:
        widths = [2100, 5700]
    else:
        widths = [7800 // cols] * cols
        widths[-1] += 7800 - sum(widths)
    set_table_geometry(table, widths)

    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 and len(value) < 28 else WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, value, size=8.5 if cols >= 4 else 9)
            if r_idx == 0:
                shade(cell._tc, PINE_LIGHT)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(PINE)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(82)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SCENICMIND  智景")
    set_run_font(r, size=13, color=AMBER, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("路演技术手册")
    set_run_font(r, size=31, color=PINE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(34)
    r = p.add_run("景区客流预测 · 经营指标分析 · 多 Agent 决策报告")
    set_run_font(r, size=14, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("从数据接入到预测、解释、审核与交付的完整技术闭环")
    set_run_font(r, size=11.5, color=INK, italic=True)

    callout = doc.add_paragraph()
    callout.alignment = WD_ALIGN_PARAGRAPH.CENTER
    callout.paragraph_format.left_indent = Inches(0.65)
    callout.paragraph_format.right_indent = Inches(0.65)
    callout.paragraph_format.space_before = Pt(34)
    callout.paragraph_format.space_after = Pt(38)
    shade(callout._p, PINE_LIGHT)
    paragraph_border(callout, side="top", color=AMBER, size="14", space="8")
    add_inline(callout, "核心技术：FlowStack 时间序列堆叠 + 8 大经营指标 + 五阶段受控 Agent 状态机", size=11)

    for text in ["版本：2026-08-30 当前代码核验版", "用途：路演 / 技术答辩 / 评委问答 / 团队交接"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        set_run_font(r, size=9.5, color=MUTED)
    doc.add_page_break()


def add_architecture_page(doc):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    p.add_run("技术架构总览")
    pic = doc.add_picture(str(DIAGRAM), width=Inches(6.75))
    pic_p = doc.paragraphs[-1]
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.space_after = Pt(5)
    doc_pr = pic._inline.docPr
    doc_pr.set("descr", "ScenicMind 多 Agent 协作架构：React、FastAPI、确定性数据工具、五阶段 Agent、LLM、Markdown 与 PDF 输出")
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(9)
    r = cap.add_run("图 1  ScenicMind 多 Agent 协作架构")
    set_run_font(r, size=8.5, color=MUTED, italic=True)

    intro = doc.add_paragraph()
    shade(intro._p, SOFT)
    paragraph_border(intro, color=AMBER)
    intro.paragraph_format.left_indent = Inches(0.12)
    intro.paragraph_format.right_indent = Inches(0.12)
    intro.paragraph_format.space_before = Pt(2)
    intro.paragraph_format.space_after = Pt(8)
    add_inline(intro, "讲图顺序：先讲上层用户与 API，再讲中层五个 Agent，最后讲底层确定性 Tool 与 LLM 分离，以及 Reviewer 返工和 LLM 降级。", size=10)
    doc.add_page_break()


def parse_markdown(doc, lines):
    idx = 0
    in_frontmatter = True
    while idx < len(lines):
        line = lines[idx].rstrip()
        stripped = line.strip()

        if in_frontmatter:
            if stripped.startswith("## 0."):
                in_frontmatter = False
            else:
                idx += 1
                continue

        if not stripped or stripped == "---":
            idx += 1
            continue

        if stripped.startswith("!["):
            idx += 1
            continue

        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline(p, stripped[4:], size=11.5, color=BLUE_DARK)
            idx += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, stripped[3:], size=17, color=PINE)
            idx += 1
            continue
        if stripped.startswith("# "):
            idx += 1
            continue

        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.right_indent = Inches(0.1)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            shade(p._p, SOFT)
            paragraph_border(p, color=AMBER)
            add_inline(p, stripped[2:], size=10.5)
            idx += 1
            continue

        if stripped.startswith("|") and idx + 1 < len(lines) and re.match(r"^\|?\s*:?-+", lines[idx + 1].strip()):
            rows = []
            idx += 2
            header = [c.strip() for c in stripped.strip("|").split("|")]
            rows.append(header)
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                rows.append([c.strip() for c in lines[idx].strip().strip("|").split("|")])
                idx += 1
            add_table(doc, rows)
            continue

        if re.match(r"^\d+\.\s+", stripped):
            num_id = create_numbering(doc, True)
            while idx < len(lines) and re.match(r"^\d+\.\s+", lines[idx].strip()):
                text = re.sub(r"^\d+\.\s+", "", lines[idx].strip())
                p = doc.add_paragraph()
                apply_numbering(p, num_id)
                add_inline(p, text)
                idx += 1
            continue

        if stripped.startswith("- "):
            num_id = create_numbering(doc, False)
            while idx < len(lines) and lines[idx].strip().startswith("- "):
                p = doc.add_paragraph()
                apply_numbering(p, num_id)
                add_inline(p, lines[idx].strip()[2:])
                idx += 1
            continue

        p = doc.add_paragraph()
        add_inline(p, stripped)
        idx += 1


def build():
    doc = Document()
    configure_section(doc.sections[0])
    configure_styles(doc)
    add_cover(doc)
    add_architecture_page(doc)
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    parse_markdown(doc, lines)

    props = doc.core_properties
    props.title = "ScenicMind 智景｜路演技术手册"
    props.subject = "景区客流预测、指标蓝图与多 Agent 决策报告平台技术说明"
    props.author = "ScenicMind Team"
    props.keywords = "ScenicMind, FlowStack, Multi-Agent, 景区客流预测, 技术路演"
    props.comments = "基于 2026-08-30 当前代码与模型产物核验生成"

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
