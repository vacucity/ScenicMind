from pathlib import Path
from zipfile import ZipFile
import json
import re

from docx import Document
from docx.oxml.ns import qn

path = Path(__file__).resolve().parent / "ScenicMind-路演技术手册.docx"
doc = Document(path)
all_text = "\n".join(p.text for p in doc.paragraphs)
headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
numbered = [p for p in doc.paragraphs if p._p.pPr is not None and p._p.pPr.numPr is not None]

issues = []
if len(doc.paragraphs) < 180:
    issues.append("paragraph_count_too_low")
if len(headings) < 30:
    issues.append("heading_count_too_low")
if len(numbered) < 20:
    issues.append("real_list_count_too_low")
if len(doc.tables) < 3:
    issues.append("table_count_too_low")
if len(doc.inline_shapes) != 1:
    issues.append("diagram_missing_or_duplicated")
if any(token in all_text for token in ("TODO", "TBD", "�")):
    issues.append("placeholder_or_bad_glyph")

table_geometry = []
for idx, table in enumerate(doc.tables):
    grid = [int(col.get(qn("w:w"))) for col in table._tbl.tblGrid]
    tbl_w_el = table._tbl.tblPr.first_child_found_in("w:tblW")
    tbl_w = int(tbl_w_el.get(qn("w:w"))) if tbl_w_el is not None else None
    widths_ok = tbl_w == sum(grid)
    cells_ok = all(
        int(cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW").get(qn("w:w"))) == grid[cidx]
        for row in table.rows for cidx, cell in enumerate(row.cells)
    )
    fixed_heights = any(row._tr.get_or_add_trPr().find(qn("w:trHeight")) is not None for row in table.rows)
    if not widths_ok or not cells_ok or fixed_heights:
        issues.append(f"table_{idx}_geometry")
    table_geometry.append({"index": idx, "columns": len(grid), "width": tbl_w, "widths_ok": widths_ok, "cells_ok": cells_ok, "fixed_heights": fixed_heights})

section = doc.sections[0]
with ZipFile(path) as archive:
    media = [n for n in archive.namelist() if n.startswith("word/media/")]
    document_xml = archive.read("word/document.xml").decode("utf-8")
    alt_ok = "ScenicMind 多 Agent 协作架构" in document_xml
    if not alt_ok:
        issues.append("image_alt_missing")

result = {
    "path": str(path),
    "size": path.stat().st_size,
    "paragraphs": len(doc.paragraphs),
    "headings": len(headings),
    "real_list_paragraphs": len(numbered),
    "tables": len(doc.tables),
    "inline_shapes": len(doc.inline_shapes),
    "media_files": media,
    "image_alt_ok": alt_ok,
    "page_inches": [round(section.page_width.inches, 2), round(section.page_height.inches, 2)],
    "margins_inches": [round(section.top_margin.inches, 2), round(section.right_margin.inches, 2), round(section.bottom_margin.inches, 2), round(section.left_margin.inches, 2)],
    "table_geometry": table_geometry,
    "issues": issues,
    "pass": not issues,
}
print(json.dumps(result, ensure_ascii=False, indent=2))
raise SystemExit(0 if not issues else 1)
